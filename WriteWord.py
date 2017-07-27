from docx import Document
import gc

import wx


class WriteWord:
    # base_document = Document()
    # base_document_path = str()
    new_document = Document()
    # new_document_path = str()

    def __init__(self, path, excel_document, progress, text):
        # Progress Bar
        self.progress = progress
        self.text = text

        # Open base word document
        self.base_document_path = path
        self.base_document = Document(path)

        # Create new document
        self.new_document_path = excel_document.get_excel_path().replace('.xlsx', '') + '_new.docx'
        self.new_document.save(self.new_document_path)

        self.path = path
        self.population = str(excel_document.get_population())

        # Make page, find and replace
        self.copying_page(excel_document.get_population())
        self.find_and_replace(excel_document.get_howmuch(), excel_document.get_numberof())

    def find_and_replace(self, excel_howmuch, excel_numberof):
        howmuch = excel_howmuch
        numberof = excel_numberof

        howmuch_counter = 0
        numberof_counter = 0
        for p in self.new_document.paragraphs:
            if "HOWMUCH" in p.text:
                inline = p.runs

                for i in range(len(inline)):
                    if "HOWMUCH" in inline[i].text:
                        text = inline[i].text.replace("HOWMUCH", str(howmuch[howmuch_counter]))
                        inline[i].text = text
                        howmuch_counter += 1

            if "NUMBEROF" in p.text:
                inline = p.runs

                for i in range(len(inline)):
                    if "NUMBEROF" in inline[i].text:
                        text = inline[i].text.replace("NUMBEROF", str(numberof[numberof_counter]).zfill(len(self.population)))
                        inline[i].text = text
                        numberof_counter += 1

            self.text.SetLabel("치환중 " + str(numberof_counter))
            self.progress.SetValue(numberof_counter)
            wx.Yield()

        self.new_document.save(self.new_document_path)

        self.text.SetLabel("완료 (아마도)")
        self.progress.SetValue(0)

    def copying_page(self, population):

        self.progress.SetRange(population)

        test = 0
        for counter in range(population):
            for p in self.base_document.element.body:
                self.new_document.element.body.append(p)
            self.base_document = Document(self.base_document_path)
            self.progress.SetValue(test)
            self.text.SetLabel("페이지 복사중 " + str(test))
            wx.Yield()
            test += 1

        self.new_document.save(self.new_document_path)
        self.progress.SetValue(0)
